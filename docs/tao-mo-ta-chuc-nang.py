"""Sinh file .docx mô tả chức năng, đọc số liệu trực tiếp từ dữ liệu trong repo.

    python docs/tao-mo-ta-chuc-nang.py

Nhờ đọc từ data/ nên các con số trong tài liệu không bị lệch khi nội dung
thay đổi — chạy lại script là tài liệu khớp lại với mã nguồn.
"""

from __future__ import annotations

import glob
import json
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "AIEDU365-Mo-ta-chuc-nang.docx"

XANH = RGBColor(0x0B, 0x4F, 0x6C)
XAM = RGBColor(0x55, 0x5F, 0x66)


def load():
    courses = [
        json.loads(Path(f).read_text(encoding="utf-8"))
        for f in sorted(glob.glob(str(ROOT / "data/courses/*.json")))
    ]
    groups = json.loads((ROOT / "data/groups.json").read_text(encoding="utf-8"))
    faqs = json.loads((ROOT / "data/faqs.json").read_text(encoding="utf-8"))
    evals = json.loads((ROOT / "data/eval/questions.json").read_text(encoding="utf-8"))
    return courses, groups, faqs, evals


def setup_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15


def h(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = XANH
        run.font.name = "Calibri"
    return p


def bullets(doc, items):
    for it in items:
        if isinstance(it, tuple):
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(it[0])
            r.bold = True
            p.add_run(" — " + it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, name in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(name)
        run.bold = True
        run.font.size = Pt(10)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    return t


def build() -> None:
    courses, groups, faqs, evals = load()
    doc = Document()
    setup_styles(doc)

    # ---------------- Trang bìa ----------------
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AIEDU365")
    r.bold = True
    r.font.size = Pt(40)
    r.font.color.rgb = XANH

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Trợ lý AI tư vấn chương trình tập huấn\nứng dụng trí tuệ nhân tạo trong giáo dục đại học")
    r.font.size = Pt(15)
    r.font.color.rgb = XAM

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("BẢN MÔ TẢ CHỨC NĂNG")
    r.bold = True
    r.font.size = Pt(17)

    for _ in range(8):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"Phạm vi: {len(courses)} khóa tập huấn {courses[0]['code']}–{courses[-1]['code']}\n"
        f"Ngày lập: {date.today().strftime('%d/%m/%Y')}"
    )
    r.font.size = Pt(11)
    r.font.color.rgb = XAM
    doc.add_page_break()

    # ---------------- 1. Giới thiệu ----------------
    h(doc, "1. Giới thiệu chung", 1)
    doc.add_paragraph(
        "AIEDU365 là hệ thống gồm một trang giới thiệu chương trình (landing page) và một "
        "trợ lý AI trả lời tự động, phục vụ công tác tuyển sinh và tư vấn cho chương trình "
        f"tập huấn ứng dụng trí tuệ nhân tạo dành cho các trường đại học, gồm {len(courses)} khóa "
        f"chuyên sâu mã {courses[0]['code']} đến {courses[-1]['code']}."
    )
    doc.add_paragraph("Hệ thống giải quyết ba việc:")
    bullets(doc, [
        ("Giới thiệu chương trình", "trình bày nội dung từng khóa một cách trực quan, thay cho việc gửi thư mời dạng văn bản."),
        ("Trả lời tự động 24/7", "chatbot giải đáp thắc mắc về nội dung, đối tượng, thời lượng, sản phẩm đầu ra của từng khóa, có dẫn nguồn."),
        ("Thu nhận và quản lý đăng ký", "ghi nhận thông tin đơn vị quan tâm, chuyển cho ban tổ chức theo dõi trong khu quản trị."),
    ])
    doc.add_paragraph(
        "Điểm khác biệt quan trọng: trợ lý AI chỉ trả lời dựa trên bộ tài liệu chính thức của "
        "chương trình. Câu hỏi nằm ngoài phạm vi sẽ được từ chối và hướng người dùng liên hệ "
        "ban tổ chức, thay vì để AI tự suy diễn ra thông tin sai."
    )

    # ---------------- 2. Tổng quan ----------------
    h(doc, "2. Tổng quan hệ thống", 1)
    doc.add_paragraph("Hệ thống chạy trên Google Cloud, gồm hai thành phần tách biệt:")
    table(doc,
          ["Thành phần", "Vai trò", "Mức truy cập"],
          [["Giao diện web", "Trang công khai và khu quản trị", "Mở ra Internet"],
           ["Dịch vụ API", "Xử lý hội thoại, truy xuất tri thức, lưu dữ liệu", "Không mở ra Internet"]])
    doc.add_paragraph()
    doc.add_paragraph(
        "Dịch vụ API cố ý không mở ra Internet. Mọi lời gọi đều phải đi qua giao diện web, "
        "nên người ngoài không thể gọi thẳng vào hệ thống xử lý dữ liệu. Đây là lựa chọn thiết kế "
        "về an toàn, không phải hạn chế kỹ thuật."
    )

    # ---------------- 3. Chức năng công khai ----------------
    h(doc, "3. Chức năng dành cho người xem", 1)

    h(doc, "3.1. Trang chủ", 2)
    bullets(doc, [
        "Giới thiệu chương trình, các con số nổi bật được tính tự động từ dữ liệu thật (số khóa, số ngày tập huấn, số module phần mềm) nên không bao giờ lệch với nội dung thực tế.",
        "Trình bày điểm khác biệt của chương trình và lộ trình triển khai.",
        "Nêu rõ nguyên tắc ứng dụng AI có trách nhiệm.",
        "Khung chat mở được ở mọi trang.",
    ])

    h(doc, "3.2. Danh mục và chi tiết khóa tập huấn", 2)
    bullets(doc, [
        f"Danh sách {len(courses)} khóa, phân theo {len(groups)} nhóm chủ đề, lọc nhanh theo nhóm.",
        "Mỗi khóa có trang riêng: mục tiêu, đối tượng, nội dung từng ngày, sản phẩm bàn giao, module phần mềm được chuyển giao, chỉ số tham chiếu, dữ liệu cần chuẩn bị, nguyên tắc AI có trách nhiệm.",
        "Có mục phân biệt với các khóa dễ nhầm lẫn, giúp đơn vị chọn đúng khóa.",
        "Công cụ so sánh hai khóa cạnh nhau.",
    ])

    h(doc, "3.3. Công cụ gợi ý chọn khóa", 2)
    doc.add_paragraph(
        "Người dùng trả lời vài câu hỏi ngắn về vị trí công tác và nhu cầu; hệ thống đề xuất "
        "khóa phù hợp nhất kèm lý do. Dành cho người chưa rõ nên bắt đầu từ đâu."
    )

    h(doc, "3.4. Hỏi đáp, phần mềm chuyển giao và đăng ký", 2)
    bullets(doc, [
        (f"Trang hỏi đáp", f"{len(faqs)} câu hỏi thường gặp đã biên soạn sẵn, tra cứu nhanh không cần hỏi chatbot."),
        ("Trang phần mềm chuyển giao", f"tổng hợp {sum(len(c.get('software', {}).get('modules', [])) for c in courses)} module phần mềm được bàn giao cho các trường."),
        ("Trang đăng ký", "biểu mẫu ghi nhận thông tin đơn vị quan tâm, chuyển thẳng vào khu quản trị."),
    ])

    # ---------------- 4. Chatbot ----------------
    h(doc, "4. Trợ lý AI tư vấn", 1)

    h(doc, "4.1. Trải nghiệm người dùng", 2)
    bullets(doc, [
        "Trả lời theo kiểu hiện dần từng chữ, người dùng không phải chờ đến khi có câu trả lời đầy đủ.",
        "Mỗi câu trả lời kèm trích dẫn nguồn — bấm vào xem được đúng phần tài liệu đã dùng để trả lời.",
        "Màn hình mở chat hiển thị ba nhóm gợi ý: câu hỏi được hỏi nhiều nhất, câu hỏi mới nhất, và câu hỏi do ban tổ chức biên soạn.",
        "Người dùng đánh giá hữu ích / chưa hữu ích cho từng câu trả lời; phản hồi này chuyển về khu quản trị để cải thiện nội dung.",
        "Khi người dùng có ý định đăng ký, chatbot chủ động hiện biểu mẫu để lại thông tin liên hệ.",
    ])

    h(doc, "4.2. Cách hệ thống bảo đảm trả lời đúng", 2)
    doc.add_paragraph(
        "Trợ lý không trả lời bằng kiến thức chung của mô hình AI, mà theo quy trình bốn bước:"
    )
    table(doc,
          ["Bước", "Nội dung"],
          [["1. Hiểu ý định", "Phân loại câu hỏi: tra cứu, so sánh, xin tư vấn chọn khóa, muốn đăng ký, hay ngoài phạm vi."],
           ["2. Tìm tài liệu", "Tìm trong bộ tri thức của chương trình, kết hợp tìm theo ngữ nghĩa và theo từ khóa, ưu tiên đúng mã khóa mà người dùng nhắc tới."],
           ["3. Soạn câu trả lời", "Mô hình AI chỉ được dùng phần tài liệu tìm được, kèm yêu cầu dẫn nguồn."],
           ["4. Kiểm duyệt đầu ra", "Rà lại câu trả lời trước khi hiển thị; phát hiện sai phạm thì chặn."]])
    doc.add_paragraph()
    doc.add_paragraph("Bốn nhóm nội dung bị chặn tuyệt đối, dù mô hình AI có sinh ra:")
    bullets(doc, [
        "Nói rằng AI tự quyết định công tác nhân sự.",
        "Nói rằng AI thay giảng viên chấm điểm.",
        "Nói rằng AI tự quyết định số phận bản thảo khoa học.",
        "Biến chỉ số tham chiếu thành cam kết kết quả.",
    ])
    doc.add_paragraph(
        "Ngoài ra, khi không tìm được tài liệu đủ liên quan, trợ lý từ chối trả lời và không kèm "
        "trích dẫn nào — tránh gợi ý sai rằng có căn cứ tài liệu."
    )

    h(doc, "4.3. Giới hạn chống lạm dụng", 2)
    table(doc,
          ["Giới hạn", "Giá trị"],
          [["Số tin nhắn mỗi phiên trong 1 giờ", "20"],
           ["Số lượt gọi mỗi địa chỉ trong 1 giờ", "100"],
           ["Độ dài tối đa một câu hỏi", "1.500 ký tự"],
           ["Số lượt để lại thông tin đăng ký mỗi địa chỉ trong 1 giờ", "10"]])

    # ---------------- 5. Khu quản trị ----------------
    doc.add_page_break()
    h(doc, "5. Khu quản trị", 1)
    doc.add_paragraph(
        "Khu quản trị gồm 12 màn hình, chia làm ba nhóm. Toàn bộ nội dung của trang công khai và "
        "của trợ lý AI đều sửa được tại đây, không cần lập trình viên can thiệp."
    )

    h(doc, "5.1. Nhóm Vận hành", 2)
    table(doc,
          ["Màn hình", "Chức năng"],
          [["Bảng điều khiển", "Thống kê tổng quan: lượt hỏi, tỷ lệ trả lời được, khóa được quan tâm nhất, số đăng ký mới."],
           ["Đăng ký", "Danh sách đơn vị để lại thông tin, cập nhật trạng thái xử lý, xuất file CSV."],
           ["Hội thoại", "Xem lại toàn bộ câu hỏi người dùng đã hỏi, lọc theo câu bị từ chối hoặc bị đánh giá chưa hữu ích; chuyển một câu hỏi hay thành mục hỏi đáp chính thức chỉ bằng một thao tác."]])

    h(doc, "5.2. Nhóm Nội dung", 2)
    table(doc,
          ["Màn hình", "Chức năng"],
          [["Khóa học", "Sửa toàn bộ thông tin từng khóa: mục tiêu, đối tượng, nội dung theo ngày, sản phẩm, module phần mềm, chỉ số. Ẩn/hiện khóa trên trang công khai."],
           ["Lịch khai giảng", "Quản lý các đợt khai giảng: thời gian, địa điểm, hình thức, trạng thái (còn chỗ, đã đầy, đã hủy)."],
           ["Hỏi đáp", "Thêm, sửa, xóa, sắp xếp các câu hỏi thường gặp."],
           ["Nội dung trang", "Sửa nội dung trang chủ: tiêu đề, giới thiệu, điểm khác biệt, lộ trình, thông tin liên hệ, lời chào và câu gợi ý của chatbot."]])
    doc.add_paragraph()
    doc.add_paragraph(
        "Mọi thay đổi nội dung đều được nạp lại vào bộ tri thức của trợ lý AI, nên chatbot trả lời "
        "theo thông tin mới ngay, không cần triển khai lại hệ thống."
    )

    h(doc, "5.3. Nhóm Hệ thống", 2)
    table(doc,
          ["Màn hình", "Chức năng"],
          [["Knowledge Base", "Xem toàn bộ các mẩu tài liệu mà trợ lý dùng để trả lời; thử một câu hỏi để xem hệ thống tìm ra tài liệu nào và điểm phù hợp bao nhiêu — dùng để kiểm tra vì sao bot trả lời chưa đúng."],
           ["Nhà cung cấp AI", "Nhập khóa API của Claude, OpenAI, Gemini; chọn bên nào dùng để trả lời, bên nào dùng để tìm kiếm; kiểm tra khóa còn hoạt động không."],
           ["Cấu hình AI", "Điều chỉnh cách trợ lý hành xử: lời nhắc hệ thống, độ dài câu trả lời, ngưỡng từ chối, số tài liệu lấy ra mỗi lần."],
           ["Người dùng", "Cấp và thu hồi quyền quản trị theo ba vai trò."],
           ["Nhật ký", "Ghi lại ai đã thay đổi gì, vào lúc nào."]])

    h(doc, "5.4. Phân quyền", 2)
    table(doc,
          ["Vai trò", "Được làm gì"],
          [["Viewer", "Chỉ xem thống kê và danh sách đăng ký."],
           ["Editor", "Thêm quyền sửa nội dung, khóa học, hỏi đáp, Knowledge Base."],
           ["Super Admin", "Toàn quyền, gồm cả quản lý người dùng và khóa API."]])

    # ---------------- 6. Nhà cung cấp AI ----------------
    h(doc, "6. Lựa chọn nhà cung cấp AI", 1)
    doc.add_paragraph(
        "Hệ thống không phụ thuộc vào một hãng AI duy nhất. Ban tổ chức tự chọn nhà cung cấp trong "
        "khu quản trị, đổi lúc nào cũng được:"
    )
    table(doc,
          ["Nhà cung cấp", "Cần nhập khóa API"],
          [["Google Vertex AI (mặc định)", "Không — dùng quyền sẵn có trên Google Cloud"],
           ["Google Gemini (AI Studio)", "Có"],
           ["Claude (Anthropic)", "Có"],
           ["OpenAI", "Có"],
           ["Chế độ thử (không gọi mạng)", "Không — dùng để kiểm thử"]])
    doc.add_paragraph()
    doc.add_paragraph("Cách xử lý khóa API:")
    bullets(doc, [
        "Khóa lưu trong Google Secret Manager, không lưu vào cơ sở dữ liệu và không ghi ra file.",
        "Nhập vào rồi thì không xem lại được, màn hình chỉ hiện 4 ký tự cuối để đối chiếu.",
        "Chỉ Super Admin được nhập và xóa khóa.",
        "Không xóa được khóa của nhà cung cấp đang sử dụng, tránh làm chatbot ngừng hoạt động.",
        "Nhật ký ghi lại việc thay đổi khóa nhưng không bao giờ ghi giá trị khóa.",
    ])

    # ---------------- 7. Nội dung ----------------
    doc.add_page_break()
    h(doc, "7. Nội dung chương trình trong hệ thống", 1)
    table(doc,
          ["Mã", "Tên khóa tập huấn", "Số ngày"],
          [[c["code"], c.get("title", ""), len(c.get("days", []))] for c in courses])
    doc.add_paragraph()
    doc.add_paragraph(f"Các khóa được xếp vào {len(groups)} nhóm chủ đề:")
    bullets(doc, [
        (g.get("name", g["id"]), g.get("description", ""))
        for g in sorted(groups, key=lambda x: x.get("order", 0))
    ])

    # ---------------- 8. Bảo mật ----------------
    h(doc, "8. An toàn thông tin", 1)
    bullets(doc, [
        ("Đăng nhập quản trị", "qua tài khoản Google/email của Firebase, không dùng mật khẩu tự quản lý."),
        ("Dịch vụ xử lý dữ liệu", "không mở ra Internet, chỉ giao diện web gọi vào được."),
        ("Khóa API", "lưu trong Secret Manager, không hiển thị lại sau khi nhập."),
        ("Xuất file danh sách đăng ký", "đã xử lý chống chèn công thức độc hại — người lạ điền vào biểu mẫu công khai không thể tấn công máy của cán bộ khi mở file Excel."),
        ("Nội dung do người dùng nhập", "được lọc trước khi hiển thị; hệ thống loại bỏ email, số điện thoại, đường dẫn và câu tự giới thiệu danh tính khỏi mục câu hỏi phổ biến."),
        ("Chống mạo nhận địa chỉ", "giới hạn lượt gọi tính theo địa chỉ do hạ tầng ghi nhận, người dùng không tự khai được."),
        ("Không có cửa hậu", "lối vào quản trị dành cho môi trường phát triển được tắt hoàn toàn khi chạy thật."),
    ])

    # ---------------- 9. Kiểm thử ----------------
    h(doc, "9. Kiểm thử và chất lượng", 1)
    doc.add_paragraph(
        f"Hệ thống có {len(evals['cases'])} câu hỏi kiểm thử chuẩn, trong đó có cả những câu "
        "bắt buộc phải bị từ chối (hỏi ngoài phạm vi, cố tình đánh lừa trợ lý). Mỗi lần sửa mã, "
        "toàn bộ bộ câu hỏi này được chạy lại tự động để bảo đảm trợ lý không trả lời sai và "
        "không bị lệch khỏi tài liệu gốc."
    )

    # ---------------- 10. Vận hành ----------------
    h(doc, "10. Vận hành", 1)
    bullets(doc, [
        "Hệ thống tự co giãn theo lượng truy cập, không có người dùng thì không tốn chi phí máy chủ.",
        "Cập nhật nội dung làm trực tiếp trên khu quản trị, có hiệu lực ngay.",
        "Triển khai lại bằng ba lệnh, có tài liệu hướng dẫn riêng viết cho người không chuyên về lập trình.",
        "Có sẵn hướng dẫn vận hành và tài liệu kiến trúc kèm theo mã nguồn.",
    ])

    doc.save(OUT)
    print(f"Đã tạo: {OUT}")


if __name__ == "__main__":
    build()
